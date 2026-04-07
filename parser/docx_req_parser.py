# ===========================================================
# FILE:         parser/docx_req_parser.py
# PROJECT:      pssgen — AI-Driven PSS + UVM + C Testbench Generator
# COPYRIGHT:    Copyright (c) 2026 BelTech Systems LLC
# LICENSE:      MIT License — see LICENSE file for details
# ===========================================================
#
# DESCRIPTION:
#   Parses a Word .docx requirements document and extracts requirement
#   statements and verification methods. Scans paragraphs for bracketed
#   requirement IDs followed by "shall" statements, and scans tables for
#   a VCRM (Verification Cross-Reference Matrix) to populate verification
#   method annotations. Returns a DocxReqResult dataclass with all
#   extracted data in document order.
#
# LAYER:        1 — parser
# PHASE:        v5a
#
# FUNCTIONS:
#   parse_docx_requirements(docx_path)
#     Read a .docx file and return a DocxReqResult with requirements,
#     verification methods, and document metadata.
#
# DEPENDENCIES:
#   Standard library:  re, dataclasses
#   Internal:          none
#   Third-party:       python-docx
#
# HISTORY:
#   v5a   2026-04-07  SB  Initial implementation; paragraph + VCRM table extraction
#
# ===========================================================
"""parser/docx_req_parser.py — Word document requirements parser.

Phase: v5a
Layer: 1 (parser)

Extracts requirement statements and verification methods from a .docx
requirements document. Paragraphs are scanned for bracketed IDs with
"shall" statements; tables are scanned for a VCRM to populate verification
method annotations.
"""
import re
from dataclasses import dataclass, field

import docx


# Matches bracketed requirement IDs at the start of a paragraph:
# [UART-PAR-001], [SYS-REQ-001], etc.
_REQ_ID_PATTERN = re.compile(
    r'^\[([A-Z][A-Z0-9]*(?:-[A-Z0-9][A-Z0-9]*)+)\]'
)

# Splits multiple methods separated by comma+space or just comma.
_METHOD_SPLIT = re.compile(r',\s*')


@dataclass
class DocxReqResult:
    """Result of parsing a Word document for requirements.

    Attributes:
        requirements: Mapping from requirement ID to requirement detail dict.
            Each dict has keys:
              "statement"     (str)       — full statement text, ID prefix stripped
              "verification"  (list[str]) — method names from VCRM, lowercased
              "waived"        (bool)      — always False at extraction time
              "waiver_reason" (str)       — always "" at extraction time
              "source"        (str)       — always "docx"
              "coverage_ref"  (str)       — VCRM Coverage Ref column, "" if absent
              "summary"       (str)       — VCRM Requirement Summary, "" if absent
        req_ids: Requirement IDs in document paragraph order.
        source_file: The docx_path as provided to the parser.
    """
    requirements: dict[str, dict] = field(default_factory=dict)
    req_ids: list[str] = field(default_factory=list)
    source_file: str = ""


def parse_docx_requirements(docx_path: str) -> DocxReqResult:
    """Read a .docx requirements document and extract requirement data.

    Scans all paragraphs for lines beginning with a bracketed requirement
    ID (e.g. ``[UART-PAR-001]``) that also contain the word "shall".
    Then scans all tables for a VCRM table (identified by "Req ID" and
    "Method" column headers) and merges verification method data into the
    requirements dict.

    Args:
        docx_path: Path to the .docx file to parse.

    Returns:
        DocxReqResult with requirements dict, ordered req_ids list, and
        the source_file path.

    Raises:
        FileNotFoundError: If docx_path does not exist.
        docx.opc.exceptions.PackageNotFoundError: If the file is not a
            valid .docx package.
    """
    doc = docx.Document(docx_path)
    result = DocxReqResult(source_file=docx_path)

    # ------------------------------------------------------------------
    # Step (a): Extract requirement paragraphs
    # ------------------------------------------------------------------
    for para in doc.paragraphs:
        text = para.text.strip()
        match = _REQ_ID_PATTERN.match(text)
        if not match:
            continue
        if "shall" not in text:
            continue

        req_id = match.group(0)   # full "[UART-PAR-001]" token
        bare_id = match.group(1)  # "UART-PAR-001"

        # Strip the bracketed ID prefix and normalize whitespace in the statement.
        statement = text[len(req_id):].strip()
        # Collapse multiple internal spaces to one.
        statement = re.sub(r'  +', ' ', statement)

        if bare_id not in result.requirements:
            result.req_ids.append(bare_id)
            result.requirements[bare_id] = {
                "statement": statement,
                "verification": [],
                "waived": False,
                "waiver_reason": "",
                "source": "docx",
                "coverage_ref": "",
                "summary": "",
            }

    # ------------------------------------------------------------------
    # Step (b): Scan tables for VCRM and merge verification methods
    # ------------------------------------------------------------------
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        # VCRM identification: header row contains both "Req ID" and "Method"
        if "Req ID" not in header_cells or "Method" not in header_cells:
            continue

        # Locate column indices
        req_id_col = header_cells.index("Req ID")
        method_col = header_cells.index("Method")
        summary_col = header_cells.index("Requirement Summary") if "Requirement Summary" in header_cells else None
        cov_ref_col = header_cells.index("Coverage Ref") if "Coverage Ref" in header_cells else None

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) <= max(req_id_col, method_col):
                continue

            vcrm_id = cells[req_id_col]
            method_text = cells[method_col]
            if not vcrm_id:
                continue

            # Parse method(s) — lowercase, split on comma
            methods: list[str] = []
            if method_text and method_text != "—":
                for m in _METHOD_SPLIT.split(method_text):
                    m = m.strip().lower()
                    if m:
                        methods.append(m)

            coverage_ref = ""
            if cov_ref_col is not None and cov_ref_col < len(cells):
                cr = cells[cov_ref_col]
                coverage_ref = "" if cr in ("—", "–", "\u2014", "\u2013", "?") else cr

            summary = ""
            if summary_col is not None and summary_col < len(cells):
                summary = cells[summary_col]

            if vcrm_id in result.requirements:
                # Update existing requirement with VCRM data
                result.requirements[vcrm_id]["verification"] = methods
                result.requirements[vcrm_id]["coverage_ref"] = coverage_ref
                result.requirements[vcrm_id]["summary"] = summary
            else:
                # VCRM-only entry (e.g. range entries like UART-REG-005–011)
                result.req_ids.append(vcrm_id)
                result.requirements[vcrm_id] = {
                    "statement": "",
                    "verification": methods,
                    "waived": False,
                    "waiver_reason": "",
                    "source": "docx",
                    "coverage_ref": coverage_ref,
                    "summary": summary,
                }

    return result
